#!/usr/bin/env python3
"""YouTube Transcriber — Blender-style UI"""

import os, re, sys, threading, tempfile, shutil, time
from pathlib import Path
from tkinter import (
    Tk, Label, Entry, Button, StringVar, OptionMenu, Listbox,
    Frame, Text, Scrollbar, messagebox, END, IntVar, Menu
)
from tkinter import ttk

from transcribe import (
    download_video, prepare_audio, transcribe_audio,
    format_output, parse_time, format_timestamp
)

THEMES = {
    "dark": {
        "bg": "#303030", "bg_header": "#383838",
        "surface": "#383838", "surface_lighter": "#424242", "surface_hover": "#4a4a4a",
        "border": "#505050",
        "text": "#e8e8e8", "text_sec": "#a0a0a0", "text_mut": "#707070",
        "accent": "#e87d0d", "accent_h": "#f09020", "accent_p": "#c06808",
        "ok": "#5fa85f", "err": "#cc4444",
        "inp": "#2a2a2a", "inp_b": "#505050",
    },
    "light": {
        "bg": "#e8e8e8", "bg_header": "#d0d0d0",
        "surface": "#ffffff", "surface_lighter": "#f0f0f0", "surface_hover": "#e0e0e0",
        "border": "#b0b0b0",
        "text": "#1a1a1a", "text_sec": "#555555", "text_mut": "#888888",
        "accent": "#e87d0d", "accent_h": "#f09020", "accent_p": "#c06808",
        "ok": "#3a8a3a", "err": "#cc3333",
        "inp": "#ffffff", "inp_b": "#b0b0b0",
    },
}

app = None


class TranscriberGUI:
    def __init__(self):
        global app
        app = self
        self.root = Tk()
        self.root.title("YouTube Transcriber")
        self.root.geometry("1100x720")
        self.root.minsize(900, 600)
        self.config_file = os.path.join(os.path.dirname(__file__), "config.txt")
        self.theme = self._load_theme()
        self.is_processing = False
        self.start_time_epoch = 0
        self.transcriptions_dir = os.path.join(os.getcwd(), "transcriptions")
        os.makedirs(self.transcriptions_dir, exist_ok=True)
        self.current_page = "transcribe"
        self._all_history = []
        self._history_files = []
        self.build_ui()
        self.refresh_history()
    
    def T(self): return THEMES[self.theme]
    
    def _load_theme(self):
        self._config = {}
        try:
            with open(self.config_file, "r") as f:
                for line in f:
                    if "=" in line:
                        k, v = line.strip().split("=", 1)
                        self._config[k] = v
        except:
            pass
        # Load save_dir from config
        if "save_dir" in self._config:
            self.transcriptions_dir = self._config["save_dir"]
            os.makedirs(self.transcriptions_dir, exist_ok=True)
        return self._config.get("theme", "dark")
    
    def _save_config(self):
        try:
            with open(self.config_file, "w") as f:
                f.write(f"theme={self.theme}\n")
                f.write(f"save_dir={self.transcriptions_dir}\n")
        except:
            pass
    
    def _save_theme(self):
        self._save_config()
    
    def build_ui(self):
        t = self.T()
        # Destroy old content
        for w in self.root.winfo_children():
            w.destroy()
        
        self.root.configure(bg=t["bg"])
        
        # Minimize to tray
        self.root.protocol("WM_DELETE_WINDOW", self._on_close)
        self.root.bind("<Unmap>", self._on_minimize)
        
        # Header
        hdr = Frame(self.root, bg=t["bg_header"])
        hdr.pack(fill="x")
        nav = Frame(hdr, bg=t["bg_header"])
        nav.pack(side="left", padx=8)
        self.nav_btns = {}
        for pid, label in [("transcribe", "Transcribe"), ("history", "History"), ("settings", "Settings")]:
            is_active = pid == self.current_page
            b = Label(nav, text=f"  {label}  ", font=("Segoe UI", 10),
                      bg=t["surface_hover"] if is_active else t["bg_header"],
                      fg=t["accent"] if is_active else t["text_sec"], cursor="hand2", padx=8, pady=8)
            b.pack(side="left")
            b.bind("<Button-1>", lambda e, p=pid: self.show_page(p))
            self.nav_btns[pid] = b
        Label(hdr, text="YouTube Transcriber", font=("Segoe UI Semibold", 10),
              bg=t["bg_header"], fg=t["text_sec"]).pack(side="right", padx=12)
        
        # Content
        self.content = Frame(self.root, bg=t["bg"])
        self.content.pack(fill="both", expand=True)
        self.content.grid_columnconfigure(0, weight=1)
        self.content.grid_rowconfigure(0, weight=1)
        
        self.pages = {}
        self.pages["transcribe"] = self._page_transcribe()
        self.pages["history"] = self._page_history()
        self.pages["settings"] = self._page_settings()
        self.show_page(self.current_page)
    
    def show_page(self, pid):
        t = self.T()
        self.current_page = pid
        for p, btn in self.nav_btns.items():
            btn.configure(bg=t["surface_hover"] if p == pid else t["bg_header"],
                          fg=t["accent"] if p == pid else t["text_sec"])
        for p, pg in self.pages.items():
            pg.grid_forget() if p != pid else pg.grid(row=0, column=0, sticky="nswe")
    
    def _entry(self, parent, **kw):
        t = self.T()
        d = dict(font=("Segoe UI", 10), bg=t["inp"], fg=t["text"], insertbackground=t["text"],
                 relief="flat", highlightthickness=1, highlightbackground=t["inp_b"], highlightcolor=t["accent"])
        d.update(kw)
        return Entry(parent, **d)
    
    def _btn(self, parent, accent=False, **kw):
        t = self.T()
        kw.setdefault("relief", "flat")
        kw.setdefault("bd", 0)
        kw.setdefault("font", ("Segoe UI", 10))
        kw.setdefault("cursor", "hand2")
        kw.setdefault("padx", 12)
        kw.setdefault("pady", 4)
        b = Button(parent, **kw)
        bg = t["accent"] if accent else t["surface_lighter"]
        fg = "#ffffff" if accent else t["text"]
        abg = t["accent_h"] if accent else t["surface_hover"]
        b.configure(bg=bg, fg=fg, activebackground=abg, activeforeground=fg)
        b.bind("<Enter>", lambda e: b.configure(bg=t["accent_h"] if accent else t["surface_hover"]))
        b.bind("<Leave>", lambda e: b.configure(bg=bg))
        return b
    
    def _card(self, parent):
        t = self.T()
        return Frame(parent, bg=t["surface"], highlightthickness=1, highlightbackground=t["border"])
    
    def _card_hdr(self, parent, title):
        t = self.T()
        h = Frame(parent, bg=t["surface_lighter"])
        h.pack(fill="x")
        Label(h, text=f"  {title}", font=("Segoe UI Bold", 9),
              bg=t["surface_lighter"], fg=t["text_sec"]).pack(side="left", padx=4, pady=4)
    
    def _time_validate(self, var):
        """Auto-format time input: adds colons, keeps cursor position."""
        def handler(*args):
            val = var.get()
            # Remove non-digits
            digits = "".join(c for c in val if c.isdigit())[:6]
            # Format as XX:XX:XX
            parts = []
            for i in range(0, len(digits), 2):
                parts.append(digits[i:i+2])
            formatted = ":".join(parts)
            if val != formatted:
                # Save cursor position relative to digits
                pos = var._last_pos if hasattr(var, '_last_pos') else len(digits)
                var.set(formatted)
                # Restore cursor: count how many real chars before cursor
                new_pos = min(pos, len(digits))
                # Convert digit index to string index (skip colons)
                str_pos = 0
                d_count = 0
                for ch in formatted:
                    if ch == ":":
                        str_pos += 1
                    else:
                        if d_count >= new_pos:
                            break
                        d_count += 1
                        str_pos += 1
        return handler
    
    def _time_key(self, var, entry):
        """Handle keypress for time input."""
        def handler(event):
            if event.keysym in ("BackSpace", "Delete"):
                val = var.get()
                digits = "".join(c for c in val if c.isdigit())
                pos = entry.index("insert")
                # Find which digit index cursor is at
                str_idx = 0
                d_idx = 0
                for i, ch in enumerate(val):
                    if i >= pos:
                        break
                    if ch != ":":
                        d_idx += 1
                if event.keysym == "BackSpace" and d_idx > 0:
                    digits = digits[:d_idx-1] + digits[d_idx:]
                elif event.keysym == "Delete" and d_idx < len(digits):
                    digits = digits[:d_idx] + digits[d_idx+1:]
                else:
                    return "break"
                parts = [digits[i:i+2] for i in range(0, len(digits), 2)]
                formatted = ":".join(parts)
                var.set(formatted)
                # Position cursor after the modified digit
                new_d_idx = d_idx - 1 if event.keysym == "BackSpace" else d_idx
                str_pos = 0
                dc = 0
                for ch in formatted:
                    if dc >= new_d_idx:
                        break
                    str_pos += 1
                    if ch != ":":
                        dc += 1
                entry.icursor(str_pos)
                return "break"
            
            if event.char and event.char.isdigit():
                val = var.get()
                digits = "".join(c for c in val if c.isdigit())
                if len(digits) >= 6:
                    return "break"
                pos = entry.index("insert")
                # Find digit index at cursor
                d_idx = 0
                for i, ch in enumerate(val):
                    if i >= pos:
                        break
                    if ch != ":":
                        d_idx += 1
                # Insert digit
                digits = digits[:d_idx] + event.char + digits[d_idx+1:] if d_idx < len(digits) else digits[:d_idx] + event.char
                parts = [digits[i:i+2] for i in range(0, len(digits), 2)]
                formatted = ":".join(parts)
                var.set(formatted)
                # Move cursor after inserted digit
                new_d_idx = d_idx + 1
                str_pos = 0
                dc = 0
                for ch in formatted:
                    if dc >= new_d_idx:
                        break
                    str_pos += 1
                    if ch != ":":
                        dc += 1
                entry.icursor(str_pos)
                return "break"
            
            # Allow Ctrl+A, arrow keys etc
            if event.state & 0x4:  # Ctrl
                return
            if event.keysym in ("Left", "Right", "Home", "End"):
                return
            return "break"
        return handler
    
    # ─── Transcribe Page ───────────────────────────────────────
    def _page_transcribe(self):
        t = self.T()
        pg = Frame(self.content, bg=t["bg"])
        pg.grid_columnconfigure(0, weight=1)
        pg.grid_rowconfigure(4, weight=1)
        
        # URL
        r0 = Frame(pg, bg=t["bg"])
        r0.grid(row=0, column=0, sticky="we", padx=12, pady=(12, 6))
        r0.grid_columnconfigure(1, weight=1)
        Label(r0, text="URL(s)", font=("Segoe UI Bold", 10), bg=t["bg"], fg=t["text"], width=6, anchor="w").grid(row=0, column=0, padx=(0, 8))
        self.url_var = StringVar()
        self.url_entry = self._entry(r0, textvariable=self.url_var)
        self.url_entry.grid(row=0, column=1, sticky="we", ipady=4)
        self.url_entry.bind("<Control-v>", self._ctrl_v)
        self.url_entry.bind("<Control-V>", self._ctrl_v)
        self._btn(r0, text="Paste", command=self._paste, width=7).grid(row=0, column=2, padx=(6, 0))
        Label(r0, text="Separate multiple URLs with comma", font=("Segoe UI", 8),
              bg=t["bg"], fg=t["text_mut"]).grid(row=1, column=1, sticky="w", pady=(2, 0))
        
        # Time + Settings
        r1 = Frame(pg, bg=t["bg"])
        r1.grid(row=1, column=0, sticky="we", padx=12, pady=6)
        
        tb = self._card(r1)
        tb.pack(side="left", fill="x", expand=True, padx=(0, 6))
        self._card_hdr(tb, "TIME RANGE")
        ti = Frame(tb, bg=t["surface"])
        ti.pack(fill="x", padx=8, pady=8)
        self.start_var = StringVar()
        self.start_entry = self._entry(ti, textvariable=self.start_var, width=10, font=("Consolas", 10), justify="center")
        self.start_entry.pack(side="left", padx=(0, 8))
        self.start_entry.bind("<Key>", self._time_key(self.start_var, self.start_entry))
        Label(ti, text="-", font=("Segoe UI", 10), bg=t["surface"], fg=t["text_mut"]).pack(side="left")
        self.end_var = StringVar()
        self.end_entry = self._entry(ti, textvariable=self.end_var, width=10, font=("Consolas", 10), justify="center")
        self.end_entry.pack(side="left", padx=(8, 8))
        self.end_entry.bind("<Key>", self._time_key(self.end_var, self.end_entry))
        Label(ti, text="HH:MM:SS", font=("Segoe UI", 8), bg=t["surface"], fg=t["text_mut"]).pack(side="left")
        
        sb = self._card(r1)
        sb.pack(side="left", fill="x", expand=True, padx=(6, 0))
        self._card_hdr(sb, "SETTINGS")
        si = Frame(sb, bg=t["surface"])
        si.pack(fill="x", padx=8, pady=8)
        
        self.lang_var = StringVar(value="auto")
        self.chunk_var = StringVar(value="5")
        self.model_var = StringVar(value="medium")
        
        for lbl, var, vals, w in [("Lang", self.lang_var, ["auto","ru","en","uk"], 6), ("Chunk", self.chunk_var, ["1","3","5","10","15"], 4), ("Model", self.model_var, ["tiny","base","small","medium","large"], 7)]:
            c = Frame(si, bg=t["surface"])
            c.pack(side="left", padx=(0, 12))
            Label(c, text=lbl, font=("Segoe UI", 8), bg=t["surface"], fg=t["text_mut"]).pack(anchor="w")
            m = OptionMenu(c, var, *vals)
            m.config(width=w, font=("Segoe UI", 10), bg=t["inp"], fg=t["text"],
                     activebackground=t["surface_hover"], activeforeground=t["text"],
                     relief="flat", highlightthickness=1, highlightbackground=t["inp_b"])
            m.pack()
        
        # Actions
        r2 = Frame(pg, bg=t["bg"])
        r2.grid(row=2, column=0, sticky="we", padx=12, pady=6)
        self._btn(r2, accent=True, text="Start Transcription", command=self._start, font=("Segoe UI Bold", 10)).pack(side="left")
        self._btn(r2, text="Stop", command=self._stop, state="disabled").pack(side="left", padx=(6, 0))
        
        pf = Frame(r2, bg=t["bg"])
        pf.pack(side="right", fill="x", expand=True, padx=(16, 0))
        self.status_var = StringVar(value="Ready")
        Label(pf, textvariable=self.status_var, font=("Segoe UI", 9), bg=t["bg"], fg=t["text_sec"]).pack(anchor="w")
        self.tl_var = StringVar(value="")
        Label(pf, textvariable=self.tl_var, font=("Segoe UI", 9), bg=t["bg"], fg=t["accent"]).pack(anchor="e")
        self.prog_var = IntVar(value=0)
        s = ttk.Style()
        s.theme_use("clam")
        s.configure("B.Horizontal.TProgressbar", background=t["accent"], troughcolor=t["surface"],
                    darkcolor=t["accent_p"], lightcolor=t["accent"], bordercolor=t["border"], thickness=8)
        ttk.Progressbar(pf, variable=self.prog_var, maximum=100, style="B.Horizontal.TProgressbar").pack(fill="x", pady=(4, 0))
        
        Frame(pg, bg=t["border"], height=1).grid(row=3, column=0, sticky="we", padx=12, pady=6)
        
        lb = self._card(pg)
        lb.grid(row=4, column=0, sticky="nswe", padx=12, pady=(0, 12))
        self._card_hdr(lb, "LOG")
        li = Frame(lb, bg=t["surface"])
        li.pack(fill="both", expand=True, padx=4, pady=4)
        self.log_text = Text(li, font=("Consolas", 9), wrap="word", relief="flat", bg=t["surface"], fg=t["text"],
                             insertbackground=t["text"], highlightthickness=0, undo=False, insertwidth=0,
                             selectbackground=t["accent"], selectforeground="#ffffff")
        self.log_text.pack(side="left", fill="both", expand=True)
        self.log_text.config(state="disabled")
        self.log_text.tag_config("success", foreground=t["ok"])
        self.log_text.tag_config("error", foreground=t["err"])
        self.log_text.tag_config("info", foreground=t["accent"])
        Scrollbar(li, command=self.log_text.yview).pack(side="right", fill="y")
        self.log_text.config(yscrollcommand=self.log_text.master.winfo_children()[-1].set)
        
        # Export section
        ef = Frame(pg, bg=t["bg"])
        ef.grid(row=5, column=0, sticky="we", padx=12, pady=(0, 12))
        self._btn(ef, text="Export as DOCX", command=self._export_docx, padx=12).pack(side="left")
        self._btn(ef, text="Export as PDF", command=self._export_pdf, padx=12).pack(side="left", padx=(8, 0))
        self.export_status = StringVar(value="")
        Label(ef, textvariable=self.export_status, font=("Segoe UI", 9), bg=t["bg"], fg=t["ok"]).pack(side="left", padx=(12, 0))
        
        return pg
    
    def _export_docx(self):
        from tkinter import filedialog
        text = self.log_text.get("1.0", END).strip()
        if not text:
            messagebox.showwarning("Warning", "No content to export")
            return
        path = filedialog.asksaveasfilename(defaultextension=".docx",
                                            filetypes=[("Word Document", "*.docx")],
                                            initialfile="transcription.docx")
        if not path:
            return
        try:
            from docx import Document
            doc = Document()
            doc.add_heading("YouTube Transcription", 0)
            for line in text.split("\n"):
                if line.startswith("["):
                    doc.add_heading(line, level=2)
                elif line.strip():
                    doc.add_paragraph(line)
            doc.save(path)
            self.export_status.set(f"Saved: {os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("Error", f"Export failed: {e}")
    
    def _export_pdf(self):
        from tkinter import filedialog
        text = self.log_text.get("1.0", END).strip()
        if not text:
            messagebox.showwarning("Warning", "No content to export")
            return
        path = filedialog.asksaveasfilename(defaultextension=".pdf",
                                            filetypes=[("PDF Document", "*.pdf")],
                                            initialfile="transcription.pdf")
        if not path:
            return
        try:
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.cell(0, 10, "YouTube Transcription", ln=True, align="C")
            pdf.ln(5)
            pdf.set_font("Arial", size=10)
            for line in text.split("\n"):
                if line.startswith("["):
                    pdf.set_font("Arial", "B", 11)
                    pdf.cell(0, 8, line, ln=True)
                    pdf.set_font("Arial", size=10)
                elif line.strip():
                    pdf.multi_cell(0, 6, line)
            pdf.output(path)
            self.export_status.set(f"Saved: {os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("Error", f"Export failed: {e}")
    
    def _hist_export_docx(self):
        from tkinter import filedialog
        text = self.hist_preview.get("1.0", END).strip()
        if not text:
            messagebox.showwarning("Warning", "No content to export")
            return
        path = filedialog.asksaveasfilename(defaultextension=".docx",
                                            filetypes=[("Word Document", "*.docx")],
                                            initialfile="transcription.docx")
        if not path:
            return
        try:
            from docx import Document
            doc = Document()
            doc.add_heading("YouTube Transcription", 0)
            for line in text.split("\n"):
                if line.startswith("["):
                    doc.add_heading(line, level=2)
                elif line.strip():
                    doc.add_paragraph(line)
            doc.save(path)
            self.hist_export_status.set(f"Saved: {os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("Error", f"Export failed: {e}")
    
    def _hist_export_pdf(self):
        from tkinter import filedialog
        text = self.hist_preview.get("1.0", END).strip()
        if not text:
            messagebox.showwarning("Warning", "No content to export")
            return
        path = filedialog.asksaveasfilename(defaultextension=".pdf",
                                            filetypes=[("PDF Document", "*.pdf")],
                                            initialfile="transcription.pdf")
        if not path:
            return
        try:
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.cell(0, 10, "YouTube Transcription", ln=True, align="C")
            pdf.ln(5)
            pdf.set_font("Arial", size=10)
            for line in text.split("\n"):
                if line.startswith("["):
                    pdf.set_font("Arial", "B", 11)
                    pdf.cell(0, 8, line, ln=True)
                    pdf.set_font("Arial", size=10)
                elif line.strip():
                    pdf.multi_cell(0, 6, line)
            pdf.output(path)
            self.hist_export_status.set(f"Saved: {os.path.basename(path)}")
        except Exception as e:
            messagebox.showerror("Error", f"Export failed: {e}")
    
    # ─── History Page ──────────────────────────────────────────
    def _page_history(self):
        t = self.T()
        pg = Frame(self.content, bg=t["bg"])
        pg.grid_columnconfigure(0, weight=2)
        pg.grid_columnconfigure(1, weight=3)
        pg.grid_rowconfigure(3, weight=1)
        
        # Header
        hdr = Frame(pg, bg=t["bg"])
        hdr.grid(row=0, column=0, columnspan=2, sticky="we", padx=12, pady=(12, 6))
        Label(hdr, text="HISTORY", font=("Segoe UI Bold", 14), bg=t["bg"], fg=t["accent"]).pack(side="left")
        self._btn(hdr, text="Downloads", command=self._open_downloads).pack(side="right", padx=(6, 0))
        self._btn(hdr, text="Transcriptions", command=self._open_transcriptions).pack(side="right")
        
        # Files search
        fs = Frame(pg, bg=t["bg"])
        fs.grid(row=1, column=0, sticky="we", padx=(12, 4), pady=6)
        fs.grid_columnconfigure(0, weight=1)
        Label(fs, text="Search files:", font=("Segoe UI", 9), bg=t["bg"], fg=t["text_sec"]).pack(anchor="w")
        self.hist_search_var = StringVar()
        self.hist_search_entry = self._entry(fs, textvariable=self.hist_search_var)
        self.hist_search_entry.pack(fill="x", ipady=3)
        self.hist_search_entry.bind("<Control-v>", self._paste_hist_search)
        self.hist_search_entry.bind("<Control-V>", self._paste_hist_search)
        self.hist_search_var.trace_add("write", lambda *a: self._filter_hist())
        
        # Preview search
        ps = Frame(pg, bg=t["bg"])
        ps.grid(row=1, column=1, sticky="we", padx=(4, 12), pady=6)
        ps.grid_columnconfigure(0, weight=1)
        Label(ps, text="Search in preview:", font=("Segoe UI", 9), bg=t["bg"], fg=t["text_sec"]).pack(anchor="w")
        search_row = Frame(ps, bg=t["bg"])
        search_row.pack(fill="x")
        search_row.grid_columnconfigure(0, weight=1)
        self.preview_search_var = StringVar()
        self.preview_search_entry = self._entry(search_row, textvariable=self.preview_search_var)
        self.preview_search_entry.grid(row=0, column=0, sticky="we", ipady=3)
        self.preview_search_entry.bind("<Return>", lambda e: self._search_preview())
        self.preview_search_entry.bind("<Control-v>", self._paste_preview_search)
        self.preview_search_entry.bind("<Control-V>", self._paste_preview_search)
        self._btn(search_row, text="Find", command=self._search_preview, padx=8).grid(row=0, column=1, padx=(6, 0))
        self._btn(search_row, text="Clear", command=self._clear_preview_search, padx=8).grid(row=0, column=2, padx=(4, 0))
        self.preview_search_status = StringVar(value="")
        Label(search_row, textvariable=self.preview_search_status, font=("Segoe UI", 8),
              bg=t["bg"], fg=t["text_mut"]).grid(row=0, column=3, padx=(8, 0))
        
        # Files list
        lb = self._card(pg)
        lb.grid(row=2, column=0, sticky="nswe", padx=(12, 4), pady=(0, 12))
        self._card_hdr(lb, "FILES")
        li = Frame(lb, bg=t["surface"])
        li.pack(fill="both", expand=True, padx=4, pady=4)
        self.hist_list = Listbox(li, font=("Segoe UI", 10), relief="flat", bd=0, bg=t["surface"], fg=t["text"],
                                 selectbackground=t["accent"], selectforeground="#ffffff", activestyle="none", highlightthickness=0)
        self.hist_list.pack(side="left", fill="both", expand=True)
        self.hist_list.bind("<<ListboxSelect>>", self._hist_select)
        Scrollbar(li, command=self.hist_list.yview).pack(side="right", fill="y")
        
        # Preview
        pb = self._card(pg)
        pb.grid(row=2, column=1, sticky="nswe", padx=(4, 12), pady=(0, 6))
        self._card_hdr(pb, "PREVIEW")
        pi = Frame(pb, bg=t["surface"])
        pi.pack(fill="both", expand=True, padx=4, pady=4)
        self.hist_preview = Text(pi, font=("Consolas", 10), wrap="word", relief="flat", bg=t["surface"], fg=t["text"],
                                 insertbackground=t["text"], highlightthickness=0, undo=False, insertwidth=0,
                                 selectbackground=t["accent"], selectforeground="#ffffff")
        self.hist_preview.tag_config("search_match", background=t["accent"], foreground="#ffffff")
        self.hist_preview.pack(side="left", fill="both", expand=True)
        self.hist_preview.bind("<Key>", lambda e: "break")
        self.hist_preview.bind("<Control-a>", lambda e: self.hist_preview.tag_add("sel", "1.0", END) or "break")
        self.hist_preview.bind("<Control-f>", lambda e: self.preview_search_entry.focus_set())
        pm = Menu(self.hist_preview, tearoff=0, bg=t["surface"], fg=t["text"], activebackground=t["accent"], activeforeground="#ffffff")
        pm.add_command(label="Copy Selected", command=self._copy_sel_hist)
        pm.add_command(label="Copy All", command=self._copy_all_hist)
        self.hist_preview.bind("<Button-3>", lambda e: pm.tk_popup(e.x_root, e.y_root))
        Scrollbar(pi, command=self.hist_preview.yview).pack(side="right", fill="y")
        
        # Export buttons for history
        ef = Frame(pg, bg=t["bg"])
        ef.grid(row=3, column=1, sticky="we", padx=(4, 12), pady=(0, 12))
        self._btn(ef, text="Export as DOCX", command=self._hist_export_docx, padx=10).pack(side="left")
        self._btn(ef, text="Export as PDF", command=self._hist_export_pdf, padx=10).pack(side="left", padx=(6, 0))
        self.hist_export_status = StringVar(value="")
        Label(ef, textvariable=self.hist_export_status, font=("Segoe UI", 9), bg=t["bg"], fg=t["ok"]).pack(side="left", padx=(10, 0))
        
        return pg
    
    def _search_preview(self):
        query = self.preview_search_var.get().strip()
        if not query:
            self.preview_search_status.set("")
            return
        
        # If same query, find next match
        if hasattr(self, '_last_search_query') and self._last_search_query == query and self._search_positions:
            self._search_idx = (self._search_idx + 1) % len(self._search_positions)
            pos = self._search_positions[self._search_idx]
            self.hist_preview.see(pos)
            self.preview_search_status.set(f"{self._search_idx + 1}/{len(self._search_positions)}")
            return
        
        # New search — find all matches
        self.hist_preview.tag_remove("search_match", "1.0", END)
        self._last_search_query = query
        self._search_positions = []
        start = "1.0"
        while True:
            pos = self.hist_preview.search(query, start, stopindex=END, nocase=True)
            if not pos:
                break
            end_pos = f"{pos}+{len(query)}c"
            self.hist_preview.tag_add("search_match", pos, end_pos)
            self._search_positions.append(pos)
            start = end_pos
        
        count = len(self._search_positions)
        self.preview_search_status.set(f"{count} found")
        if count > 0:
            self._search_idx = 0
            self.hist_preview.see(self._search_positions[0])
    
    def _clear_preview_search(self):
        self.preview_search_var.set("")
        self.preview_search_status.set("")
        self.hist_preview.tag_remove("search_match", "1.0", END)
        self._search_positions = []
        self._search_idx = 0
        self._last_search_query = ""
        self.hist_preview.tag_remove("search_match", "1.0", END)
    
    # ─── Settings Page ─────────────────────────────────────────
    def _page_settings(self):
        t = self.T()
        pg = Frame(self.content, bg=t["bg"])
        
        hdr = Frame(pg, bg=t["bg"])
        hdr.pack(fill="x", padx=12, pady=(12, 6))
        Label(hdr, text="SETTINGS", font=("Segoe UI Bold", 14), bg=t["bg"], fg=t["accent"]).pack(side="left")
        
        # Appearance
        c = self._card(pg)
        c.pack(fill="x", padx=12, pady=6)
        self._card_hdr(c, "APPEARANCE")
        ci = Frame(c, bg=t["surface"])
        ci.pack(fill="x", padx=8, pady=12)
        Label(ci, text="Theme", font=("Segoe UI", 10), bg=t["surface"], fg=t["text"]).pack(side="left")
        self._btn(ci, text="Light", command=lambda: self._set_theme("light")).pack(side="right", padx=(8, 0))
        self._btn(ci, text="Dark", command=lambda: self._set_theme("dark")).pack(side="right")
        
        # Storage
        s = self._card(pg)
        s.pack(fill="x", padx=12, pady=6)
        self._card_hdr(s, "STORAGE")
        si = Frame(s, bg=t["surface"])
        si.pack(fill="x", padx=8, pady=12)
        Label(si, text="Save folder:", font=("Segoe UI", 10), bg=t["surface"], fg=t["text"]).pack(side="left")
        self.save_dir_var = StringVar(value=self.transcriptions_dir)
        save_entry = self._entry(si, textvariable=self.save_dir_var, font=("Consolas", 9))
        save_entry.pack(side="left", fill="x", expand=True, padx=(8, 8))
        self._btn(si, text="Browse", command=self._browse_save_dir, padx=8).pack(side="left")
        self._btn(si, text="Reset", command=self._reset_save_dir, padx=8).pack(side="left", padx=(4, 0))
        
        # About
        a = self._card(pg)
        a.pack(fill="x", padx=12, pady=6)
        self._card_hdr(a, "ABOUT")
        ai = Frame(a, bg=t["surface"])
        ai.pack(fill="x", padx=8, pady=8)
        Label(ai, text="YouTube Transcriber v1.1", font=("Segoe UI", 10), bg=t["surface"], fg=t["text"]).pack(anchor="w")
        Label(ai, text="github.com/viketvova/youtube-transcriber", font=("Segoe UI", 9), bg=t["surface"], fg=t["accent"]).pack(anchor="w", pady=(2, 0))
        
        return pg
    
    def _browse_save_dir(self):
        from tkinter import filedialog
        d = filedialog.askdirectory(initialdir=self.transcriptions_dir)
        if d:
            self.save_dir_var.set(d)
            self.transcriptions_dir = d
            self._save_config()
            self.refresh_history()
    
    def _reset_save_dir(self):
        default = os.path.join(os.getcwd(), "transcriptions")
        self.save_dir_var.set(default)
        self.transcriptions_dir = default
        os.makedirs(default, exist_ok=True)
        self._save_config()
        self.refresh_history()
    
    # ─── Theme ─────────────────────────────────────────────────
    def _set_theme(self, theme):
        self.theme = theme
        self._save_theme()
        self.build_ui()
        self.refresh_history()
    
    # ─── Clipboard ─────────────────────────────────────────────
    def _ctrl_v(self, e=None):
        try:
            self.url_entry.delete(0, END)
            self.url_entry.insert(0, self.root.clipboard_get().strip())
        except: pass
        return "break"
    
    def _paste_to_entry(self, entry):
        try:
            text = self.root.clipboard_get().strip()
            entry.delete(0, END)
            entry.insert(0, text)
        except: pass
    
    def _paste_preview_search(self, event=None):
        try:
            text = self.root.clipboard_get().strip()
            self.preview_search_entry.delete(0, END)
            self.preview_search_entry.insert(0, text)
        except: pass
        return "break"
    
    def _paste_hist_search(self, event=None):
        try:
            text = self.root.clipboard_get().strip()
            self.hist_search_entry.delete(0, END)
            self.hist_search_entry.insert(0, text)
        except: pass
        return "break"
    
    def _paste(self):
        try: self.url_var.set(self.root.clipboard_get().strip())
        except: pass
    
    def _on_close(self):
        self.root.destroy()
    
    def _on_minimize(self, event=None):
        if self.root.state() == "iconic":
            self.root.after(100, lambda: self.root.state("iconic"))
    
    def _copy_all_hist(self):
        self.root.clipboard_clear()
        self.root.clipboard_append(self.hist_preview.get("1.0", END).strip())
    
    def _copy_sel_hist(self):
        try:
            self.root.clipboard_clear()
            self.root.clipboard_append(self.hist_preview.get("sel.first", "sel.last"))
        except: pass
    
    # ─── Log ───────────────────────────────────────────────────
    def log(self, msg, tag=None):
        self.log_text.config(state="normal")
        self.log_text.insert(END, msg + "\n", tag) if tag else self.log_text.insert(END, msg + "\n")
        self.log_text.see(END)
        self.log_text.config(state="disabled")
    
    def clear_log(self):
        self.log_text.config(state="normal")
        self.log_text.delete("1.0", END)
        self.log_text.config(state="disabled")
    
    # ─── History ───────────────────────────────────────────────
    def refresh_history(self):
        self._all_history = []
        self.hist_list.delete(0, END)
        if not os.path.isdir(self.transcriptions_dir):
            return
        for f in sorted(Path(self.transcriptions_dir).glob("*.txt"), key=os.path.getmtime, reverse=True):
            self._all_history.append((f.stem, str(f)))
        self._filter_hist()
    
    def _filter_hist(self):
        q = self.hist_search_var.get().strip().lower()
        self.hist_list.delete(0, END)
        self._history_files = []
        for name, path in self._all_history:
            if q and q not in name.lower(): continue
            self.hist_list.insert(END, name)
            self._history_files.append(path)
    
    def _hist_select(self, e=None):
        sel = self.hist_list.curselection()
        if not sel: return
        i = sel[0]
        if i < len(self._history_files):
            try:
                with open(self._history_files[i], "r", encoding="utf-8") as f: c = f.read()
                self.hist_preview.delete("1.0", END)
                self.hist_preview.insert("1.0", c)
            except Exception as e:
                self.hist_preview.delete("1.0", END)
                self.hist_preview.insert("1.0", f"Error: {e}")
    
    def _open_transcriptions(self): os.startfile(self.transcriptions_dir)
    def _open_downloads(self):
        d = os.path.join(os.path.dirname(__file__), "downloads")
        os.makedirs(d, exist_ok=True)
        os.startfile(d)
    
    # ─── Progress ──────────────────────────────────────────────
    def _set_progress(self, v, s=None, tl=None):
        self.prog_var.set(v)
        if s: self.status_var.set(s)
        if tl is not None: self.tl_var.set(tl)
        # Update window title with progress
        if v > 0 and v < 100:
            self.root.title(f"[{int(v)}%] YouTube Transcriber")
        elif v >= 100:
            self.root.title("YouTube Transcriber - Done!")
        else:
            self.root.title("YouTube Transcriber")
        self.root.update_idletasks()
    
    def validate_url(self, url): return bool(re.match(r'https?://(www\.)?(youtube\.com|youtu\.be)', url))
    def _stop(self):
        self.is_processing = False
        self.log("Stopping...", "info")
    
    # ─── Transcription ─────────────────────────────────────────
    def _start(self):
        url_input = self.url_var.get().strip()
        if not url_input:
            messagebox.showwarning("Warning", "Please enter a YouTube URL")
            return
        
        # Parse multiple URLs (comma or newline separated)
        urls = [u.strip() for u in url_input.replace("\n", ",").split(",") if u.strip()]
        
        # Validate all URLs
        for url in urls:
            if not self.validate_url(url):
                messagebox.showerror("Error", f"Invalid YouTube URL:\n{url}")
                return
        
        self.clear_log()
        self._set_progress(0, "Starting...", "")
        self.is_processing = True
        self.start_time_epoch = time.time()
        
        if len(urls) == 1:
            threading.Thread(target=self._run, args=(urls[0],), daemon=True).start()
        else:
            self.log(f"Batch mode: {len(urls)} videos to process", "info")
            threading.Thread(target=self._run_batch, args=(urls,), daemon=True).start()
    
    def _run_batch(self, urls):
        total = len(urls)
        for i, url in enumerate(urls, 1):
            if not self.is_processing:
                break
            self.log(f"\n{'='*60}", "info")
            self.log(f"[{i}/{total}] Processing: {url}", "info")
            self.log(f"{'='*60}", "info")
            self._run(url, batch=True)
        self.is_processing = False
        self.root.title("YouTube Transcriber")
        if self.is_processing == False:
            self.log(f"\nBatch complete: {total} videos processed", "success")
    
    def _update_tl(self, p):
        if p > 0:
            el = time.time() - self.start_time_epoch
            rem = el / (p / 100) - el
            self.tl_var.set(f"~{int(rem//60)}m {int(rem%60)}s left" if rem > 0 else "")
    
    def _run(self, url, batch=False):
        lang = self.lang_var.get()
        model = self.model_var.get()
        chunk = int(self.chunk_var.get())
        ss = parse_time(self.start_var.get().strip()) if self.start_var.get().strip() else 0.0
        es = parse_time(self.end_var.get().strip()) if self.end_var.get().strip() else 0.0
        
        ln = {"auto": "Auto-detect", "ru": "Russian", "en": "English", "uk": "Ukrainian"}
        self.log(f"URL: {url}", "info")
        self.log(f"Language: {ln.get(lang, lang)} | Model: {model} | Chunk: {chunk}min", "info")
        if ss > 0 or es > 0:
            self.log(f"Time: {format_timestamp(ss)} - {format_timestamp(es) if es > 0 else 'end'}", "info")
        self.log("-" * 60)
        
        dl = os.path.join(os.path.dirname(__file__), "downloads")
        os.makedirs(dl, exist_ok=True)
        td = tempfile.mkdtemp(prefix="yt_", dir=dl)
        
        try:
            self._set_progress(5, "Downloading...", "")
            self.log("[1/3] Downloading video...")
            try: vp = download_video(url, td)
            except Exception as e:
                self.log(f"Download error: {e}", "error"); return
            if not vp:
                self.log("Download failed", "error"); return
            if not self.is_processing: return
            
            self._set_progress(35, "Preparing audio...", "")
            self.log("[2/3] Preparing audio...")
            ap = prepare_audio(vp, td, ss, es)
            if not ap:
                self.log("Audio preparation failed", "error"); return
            if not self.is_processing: return
            
            self._set_progress(45, "Loading model...", "")
            self.log(f"[3/3] Transcribing ({model})...")
            
            def pcb(d, _):
                p = min(45 + int(d * 0.5), 95)
                self._set_progress(p, f"Segments: {d}", "")
                self._update_tl(p)
            
            segs = transcribe_audio(ap, model_size=model, language=lang, time_offset=ss, progress_callback=pcb)
            if not segs:
                self.log("Transcription failed", "error"); return
            
            self._set_progress(98, "Saving...", "")
            txt = format_output(segs, url, chunk)
            
            try:
                import subprocess, json
                r = subprocess.run([sys.executable, "-m", "yt_dlp", "--dump-json", "--skip-download", url],
                                   capture_output=True, text=True, timeout=30)
                title = json.loads(r.stdout).get("title", "") if r.returncode == 0 else ""
            except: title = ""
            
            vid_match = re.search(r'(?:v=|youtu\.be/)([^&?]+)', url)
            if title:
                fn = re.sub(r'[<>:"/\\|?*]', '_', title)[:80] + ".txt"
            elif vid_match:
                fn = f"transcript_{vid_match.group(1)[:11]}.txt"
            else:
                fn = "transcript.txt"
            
            op = os.path.join(self.transcriptions_dir, fn)
            with open(op, "w", encoding="utf-8") as f: f.write(txt)
            
            el = time.time() - self.start_time_epoch
            self._set_progress(100, "Done!", f"{int(el//60)}m {int(el%60)}s")
            self.log("-" * 60)
            self.log(f"Saved: {fn}", "success")
            self.log(f"Segments: {len(segs)} | Time: {int(el//60)}m {int(el%60)}s", "success")
            self.refresh_history()
        except Exception as e:
            self.log(f"Error: {e}", "error")
            self._set_progress(0, "Error", "")
        finally:
            self.is_processing = False
    
    def run(self): self.root.mainloop()


if __name__ == "__main__":
    try:
        TranscriberGUI().run()
    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        print(tb)
        with open(os.path.join(os.path.dirname(__file__), "error.log"), "w") as f: f.write(tb)
        input("Press Enter...")
